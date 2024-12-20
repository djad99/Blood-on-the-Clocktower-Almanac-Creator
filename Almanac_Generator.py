import json

import docx
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION

from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"

file_name = ""

characters = []

script_file = ""
master_almanac_filename = "almanac.json"
master_almanac = json.load(open(master_almanac_filename, "r"))

def init_characters():
    global master_almanac

    for i in master_almanac["characters"]:
        characters.append(i)


# Copied and pasted from https://github.com/python-openxml/python-docx/issues/105#issuecomment-442786431
def add_horizontal_line(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def add_characters(doc, is_script_almanac):
    # This creates 2 columns -- seems api support is not coming outside of this
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, "2")

    first_character = True
    count = 0
    is_even = True
    cached_type = "Townsfolk"
    for character in characters:
        character_information = master_almanac["characters"][character]
        character_color = RGBColor(0x06, 0x63, 0xB9)
        if character_information[8]["alignment"] == "Evil":
            character_color = RGBColor(0x9B, 0x00, 0x00)

        # If creating a script almanac, the character type changes, and there's an uneven amount of characters
        # in a section, make a column break
        if character_information[10]["type"] != cached_type and not is_even and is_script_almanac:
            dummy_paragraph = doc.add_paragraph()
            dummy_paragraph_run = dummy_paragraph.add_run()
            dummy_paragraph_run.add_break(WD_BREAK.COLUMN)
            is_even = True

        image_file_path = "Images/Official Characters/"

        # Add a centered picture with height of 1 inch
        if not first_character:
            picture_run = doc.paragraphs[-1].add_run()
            picture_run.add_picture(image_file_path+character_information[9]["image_file_name"], height=Inches(1))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            first_character = False
            doc.add_picture(image_file_path+character_information[9]["image_file_name"], height=Inches(1))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        character_name = doc.add_paragraph()
        character_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        character_name_runner = character_name.add_run(character_information[0]["name"])
        character_name_runner.bold = True
        character_name_runner.font.size = Pt(22)
        character_name_runner.font.name = "Dumbledor 1"
        character_name_runner.font.color.rgb = character_color
        character_name.line_spacing_rule = WD_LINE_SPACING.SINGLE

        character_ability_runner = character_name.add_run("\n"+character_information[1]["ability"])
        character_ability_runner.font.size = Pt(9)
        character_ability_runner.font.name = "Franklin Gothic Book"
        character_ability_runner.font.color.rgb = character_color

        add_horizontal_line(character_name)

        character_quote = doc.add_paragraph()
        character_quote.paragraph_format.line_spacing = 1
        character_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
        character_quote_runner = character_quote.add_run('"' + character_information[2]["quote"] + '"')
        character_quote_runner.italic = True
        character_quote_runner.font.size = Pt(9)
        character_quote_runner.font.name = "Trade Gothic LT Std"
        character_quote_runner.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        character_summary = doc.add_paragraph()
        character_summary.paragraph_format.line_spacing = 1
        character_summary_runner = character_summary.add_run(character_information[3]["summary"])
        character_summary_runner.font.size = Pt(9)
        character_summary_runner.font.name = "Franklin Gothic Book"

        bullets = character_information[4]["bullets"]

        for bullet in bullets:
            item = doc.add_paragraph()
            item.paragraph_format.line_spacing = 1
            item.style = doc.styles['List Bullet 2']
            item_runner = item.add_run(bullet)
            item_runner.font.size = Pt(9)
            item_runner.font.name = "Franklin Gothic Book"

        examples_header = doc.add_paragraph()
        examples_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        examples_head_run = examples_header.add_run("EXAMPLES")
        examples_head_run.font.size = Pt(12)
        examples_head_run.bold = True
        examples_head_run.font.name = "Dumbledor 1"
        examples_head_run.font.color.rgb = character_color

        examples = character_information[5]["examples"]

        for example in examples:
            item = doc.add_paragraph()
            item.paragraph_format.line_spacing = 1
            item_runner = item.add_run(example)
            item_runner.font.size = Pt(9)
            item_runner.font.name = "Franklin Gothic Book"

        how_to_run_header = doc.add_paragraph()
        how_to_run_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        how_to_run_head_run = how_to_run_header.add_run("HOW TO RUN")
        how_to_run_head_run.font.size = Pt(12)
        how_to_run_head_run.bold = True
        how_to_run_head_run.font.name = "Dumbledor 1"
        how_to_run_head_run.font.color.rgb = character_color

        how_to_run_paragraphs = character_information[6]["how-to-run"]
        for paragraph in how_to_run_paragraphs:
            how_to_run = doc.add_paragraph()
            how_to_run.paragraph_format.line_spacing = 1
            how_to_run_run = how_to_run.add_run(paragraph)
            how_to_run_run.font.size = Pt(9)
            how_to_run_run.font.name = "Franklin Gothic Book"

        boxes = character_information[7]["boxes"]
        if len(boxes) > 0:
            extra_info = doc.add_paragraph()
            extra_info.paragraph_format.line_spacing = 1
            # Shading the boxes
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D3D3D3')
            extra_info.paragraph_format.element.get_or_add_pPr()
            extra_info.paragraph_format.element.pPr.append(shd)

            box_count = 1
            for box in boxes:
                if box_count < len(boxes):
                    extra_info_run = extra_info.add_run(box + "\n")
                else:
                    extra_info_run = extra_info.add_run(box)
                box_count += 1
                extra_info_run.font.size = Pt(9)
                extra_info_run.font.italic = True
                extra_info_run.font.name = "Franklin Gothic Book"

        count += 1
        is_even = not is_even
        cached_type = character_information[10]["type"]
        if count != len(characters):
            dummy_paragraph = doc.add_paragraph()
            dummy_paragraph_run = dummy_paragraph.add_run()
            dummy_paragraph_run.add_break(WD_BREAK.COLUMN)

    return doc


def filter_characters():
    global characters
    script_file = input("Enter the script's file name: ")
    script_characters = json.load(open("Scripts/"+script_file))
    characters_appearing = []

    first = True

    for i in script_characters:
        if not first:
            characters_appearing.append(i)
        else:
            first = False

    tmp_list = [value for value in characters_appearing if value in characters]

    # Got all of the characters we need, now to sort them by character type.

    townsfolk = []
    outsiders = []
    minions = []
    demons = []

    for character in tmp_list:
        if master_almanac["characters"][character][10]["type"] == "Townsfolk":
            townsfolk.append(character)
        elif master_almanac["characters"][character][10]["type"] == "Outsider":
            outsiders.append(character)
        elif master_almanac["characters"][character][10]["type"] == "Minion":
            minions.append(character)
        else:
            demons.append(character)

    characters = townsfolk
    characters.extend(outsiders)
    characters.extend(minions)
    characters.extend(demons)


def create_master_almanac():
    characters.sort()
    # If creating a master almanac, all characters are in alphabetical order
    doc = docx.Document()

    title = "All Characters Almanac"
    title_page = doc.add_paragraph()
    title_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_page_runner = title_page.add_run("\n\n\n\n\n\n\n\n" + title)
    title_page_runner.bold = True
    title_page_runner.font.size = Pt(30)
    title_page_runner.font.name = "Dumbledor 1"

    title_page_addendum = title_page.add_run("\nUnofficial Almanac for Blood on the Clocktower")
    title_page_addendum.font.size = Pt(20)
    title_page_addendum.font.name = "Dumbledor 1"

    doc.add_page_break()

    doc = add_characters(doc, False)

    doc.save("Master Almanac.docx")


def intro_paragraph_page(doc):
    intro_filename = input("Filename for script introduction (leave blank for none): ")
    if intro_filename == "":
        return doc

    intro_file = open("Introductions/"+intro_filename)

    lines = intro_file.readlines()

    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_run = p.add_run(line)
        p_run.bold = True
        p_run.font.size = Pt(16)
        p_run.font.name = "Dumbledor 1"
    return doc



def create_sub_almanac():
    doc = docx.Document()

    title = input("Enter the name of the script: ")
    title_page = doc.add_paragraph()
    title_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_page_runner = title_page.add_run("\n\n\n\n\n\n\n\n" + title)
    title_page_runner.bold = True
    title_page_runner.font.size = Pt(30)
    title_page_runner.font.name = "Dumbledor 1"

    title_page_addendum = title_page.add_run("\nUnofficial Almanac for Blood on the Clocktower")
    title_page_addendum.font.size = Pt(20)
    title_page_addendum.font.name = "Dumbledor 1"

    doc.add_page_break()

    doc = intro_paragraph_page(doc)

    filter_characters()

    doc = add_characters(doc, True)

    doc.save(title+".docx")


def transform_character_name(character_name):
    character_name = character_name.lower()
    character_name = character_name.replace(" ", "")
    character_name = character_name.replace("'", "")
    character_name = character_name.replace("-", "")
    return character_name


def add_character():
    name = input("New Character Name : ")
    ability = input("New Character Ability : ")
    quote = input("New Character Quote : ")
    summary = input("New Character Summary : ")
    bullets = []
    bullet_number = int(input("How many bullet points are there: "))
    for i in range(bullet_number):
        bullets.append(input("Bullet : "))
    examples = []
    example_count = int(input("How many examples are there: "))
    for i in range(example_count):
        examples.append(input("Example : "))
    how_to_run = []
    how_to_run_count = int(input("How many paragraphs in the how to run section are there: "))
    for i in range(how_to_run_count):
        how_to_run.append(input("How to run paragraph : "))
    boxes = []
    box_count = int(input("How many boxes are there: "))
    for i in range(box_count):
        boxes.append(input("Box : "))
    icon_file_path = "Icon_"+transform_character_name(name)+".png"
    type = input("Character type: ")

    alignment = "Evil"
    if type == "Townsfolk" or type == "Outsider":
        alignment = "Good"

    addition = [{
        "name": name.title()
    }, {
        "ability": ability
    }, {
        "quote": quote
    }, {
        "summary": summary
    }, {
        "bullets": bullets
    }, {
        "examples": examples
    }, {
        "how-to-run": how_to_run
    }, {
        "boxes": boxes
    }, {
        "alignment": alignment
    }, {
        "image_file_name": icon_file_path
    }, {
        "type": type
    }
    ]

    master_almanac["characters"].update(
        {transform_character_name(name): addition}
    )

    with open(master_almanac_filename, 'w') as f:
        json.dump(master_almanac, f, indent=4)

init_characters()
if "y" in input("Would you like to add a character: ").lower():
    add_character()
elif "y" in input("Would you want to create a master almanac: ").lower():
    create_master_almanac()
else:
    create_sub_almanac()